"""
Unit tests for utils/file_handler.py
"""

import io

import pytest
from docx import Document

from utils.file_handler import (
    extract_text,
    extract_from_txt,
    extract_from_docx,
)


# ---------------------------------------------------------------------
# TXT
# ---------------------------------------------------------------------

def test_extract_txt():

    data = b"Hello SmartRFP"

    text = extract_text("sample.txt", data)

    assert text == "Hello SmartRFP"


def test_txt_latin1():

    text = "Résumé"

    data = text.encode("latin-1")

    result = extract_from_txt(data)

    assert result == text


# ---------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------

def test_extract_docx():

    doc = Document()

    doc.add_paragraph("Security Requirements")

    doc.add_paragraph("Vendor shall support SSO.")

    stream = io.BytesIO()

    doc.save(stream)

    text = extract_text(
        "sample.docx",
        stream.getvalue(),
    )

    assert "Security Requirements" in text

    assert "Vendor shall support SSO." in text


def test_extract_docx_table():

    doc = Document()

    table = doc.add_table(rows=2, cols=2)

    table.cell(0, 0).text = "Requirement"

    table.cell(0, 1).text = "Support"

    table.cell(1, 0).text = "SSO"

    table.cell(1, 1).text = "Yes"

    stream = io.BytesIO()

    doc.save(stream)

    text = extract_text(
        "table.docx",
        stream.getvalue(),
    )

    assert "Requirement" in text

    assert "Support" in text

    assert "SSO" in text

    assert "Yes" in text


# ---------------------------------------------------------------------
# Cleaning
# ---------------------------------------------------------------------

def test_cleaning_removes_extra_blank_lines():

    raw = b"""

Security


Vendor shall support SSO.




Vendor shall encrypt data.



"""

    text = extract_text("sample.txt", raw)

    assert "Security" in text

    assert "Vendor shall support SSO." in text

    assert "Vendor shall encrypt data." in text

    assert "\n\n\n" not in text


def test_cleaning_removes_multiple_spaces():

    raw = b"Vendor      shall       support     SSO."

    text = extract_text("sample.txt", raw)

    assert text == "Vendor shall support SSO."


# ---------------------------------------------------------------------
# Errors
# ---------------------------------------------------------------------

def test_unsupported_extension():

    with pytest.raises(ValueError):

        extract_text(
            "image.png",
            b"abc",
        )


def test_legacy_doc_not_supported():

    with pytest.raises(ValueError):

        extract_text(
            "legacy.doc",
            b"abc",
        )


def test_empty_text_file():

    with pytest.raises(ValueError):

        extract_text(
            "empty.txt",
            b"",
        )


def test_whitespace_only_file():

    with pytest.raises(ValueError):

        extract_text(
            "blank.txt",
            b"      \n\n\n",
        )


# ---------------------------------------------------------------------
# Mixed content
# ---------------------------------------------------------------------

def test_multiline_text():

    raw = b"""
Section 1

Vendor shall support SSO.

Vendor must encrypt data.

Vendor should provide migration services.
"""

    result = extract_text(
        "rfp.txt",
        raw,
    )

    assert "Section 1" in result

    assert "Vendor shall support SSO." in result

    assert "Vendor must encrypt data." in result

    assert "Vendor should provide migration services." in result


def test_unicode_txt():

    text = "Encryption ✓ Security"

    result = extract_text(
        "unicode.txt",
        text.encode("utf-8"),
    )

    assert "Encryption" in result

    assert "Security" in result