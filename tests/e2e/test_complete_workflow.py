"""
End-to-end tests.

These execute the complete SmartRFP workflow:

File
    ↓
Text extraction
    ↓
Requirement extraction
    ↓
Pricing
    ↓
RAG
    ↓
Draft generation
    ↓
Database persistence
    ↓
Export

Only external services (LLM / Tavily) are mocked.
"""

import io

from docx import Document

import database

from pipeline import run_pipeline
from utils.file_handler import extract_text
from utils.exporter import (
    export_docx,
    export_pdf,
    export_txt,
)


# --------------------------------------------------------------------
# Helpers
# --------------------------------------------------------------------

def make_docx():

    doc = Document()

    doc.add_heading("Security")

    doc.add_paragraph(
        "Vendor shall support MFA."
    )

    buffer = io.BytesIO()

    doc.save(buffer)

    return buffer.getvalue()


def fake_chat(*args, **kwargs):

    return """
[
 {
  "section":"Security",
  "text":"Vendor shall support MFA."
 }
]
"""


# --------------------------------------------------------------------
# Test 1
# --------------------------------------------------------------------

def test_complete_workflow(monkeypatch, temp_db):

    monkeypatch.setattr(
        "llm.chat",
        fake_chat,
    )

    database.add_kb_doc(
        "Security Guide",
        "Guide",
        "Supports MFA and SSO."
    )

    data = make_docx()

    text = extract_text(
        "rfp.docx",
        data,
    )

    rfp = database.create_rfp(
        "Deal",
        "Client",
        "UK",
        "",
        "",
        "",
        "rfp.docx",
        text,
        "",
        "",
        False,
    )

    result = run_pipeline(
        rfp,
        text,
        use_web_search=False,
    )

    assert result["requirements"] > 0

    assert result["sections"] > 0

    assert database.get_rfp(rfp) is not None

    assert len(database.get_requirements(rfp)) > 0

    assert len(database.get_draft_sections(rfp)) > 0


# --------------------------------------------------------------------
# Test 2
# --------------------------------------------------------------------

def test_docx_export(monkeypatch, temp_db):

    monkeypatch.setattr(
        "llm.chat",
        fake_chat,
    )

    database.add_kb_doc(
        "Security Guide",
        "Guide",
        "Supports MFA."
    )

    text = "Vendor shall support MFA."

    rfp = database.create_rfp(
        "Deal",
        "Client",
        "",
        "",
        "",
        "",
        "rfp.txt",
        text,
        "",
        "",
        False,
    )

    run_pipeline(
        rfp,
        text,
        use_web_search=False,
    )

    output = export_docx(rfp)

    assert len(output) > 100


# --------------------------------------------------------------------
# Test 3
# --------------------------------------------------------------------

def test_pdf_export(monkeypatch, temp_db):

    monkeypatch.setattr(
        "llm.chat",
        fake_chat,
    )

    database.add_kb_doc(
        "Security Guide",
        "Guide",
        "Supports MFA."
    )

    text = "Vendor shall support MFA."

    rfp = database.create_rfp(
        "Deal",
        "Client",
        "",
        "",
        "",
        "",
        "rfp.txt",
        text,
        "",
        "",
        False,
    )

    run_pipeline(
        rfp,
        text,
    )

    pdf = export_pdf(rfp)

    assert len(pdf) > 100


# --------------------------------------------------------------------
# Test 4
# --------------------------------------------------------------------

def test_txt_export(monkeypatch, temp_db):

    monkeypatch.setattr(
        "llm.chat",
        fake_chat,
    )

    database.add_kb_doc(
        "Security Guide",
        "Guide",
        "Supports MFA."
    )

    text = "Vendor shall support MFA."

    rfp = database.create_rfp(
        "Deal",
        "Client",
        "",
        "",
        "",
        "",
        "rfp.txt",
        text,
        "",
        "",
        False,
    )

    run_pipeline(
        rfp,
        text,
    )

    txt = export_txt(rfp)

    assert len(txt) > 100


# --------------------------------------------------------------------
# Test 5
# --------------------------------------------------------------------

def test_database_after_pipeline(monkeypatch, temp_db):

    monkeypatch.setattr(
        "llm.chat",
        fake_chat,
    )

    database.add_kb_doc(
        "Security Guide",
        "Guide",
        "Supports MFA."
    )

    text = "Vendor shall support MFA."

    rfp = database.create_rfp(
        "Deal",
        "Client",
        "",
        "",
        "",
        "",
        "rfp.txt",
        text,
        "",
        "",
        False,
    )

    run_pipeline(
        rfp,
        text,
    )

    assert database.get_pricing(rfp)

    assert database.get_draft_sections(rfp)

    assert database.get_requirements(rfp)

    assert database.get_audit_log(rfp)

    assert database.get_evaluation_metrics(rfp)